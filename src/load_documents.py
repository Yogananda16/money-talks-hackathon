import os
from docx import Document
import openpyxl

DOCS_DIR = "../documents"

def read_docx(path):
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            text += "\n" + " | ".join(cells)
    return text

def read_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            if any(row):
                text += " | ".join(str(c) for c in row if c is not None) + "\n"
    return text

def load_all_documents():
    all_text = {}
    for fname in os.listdir(DOCS_DIR):
        path = os.path.join(DOCS_DIR, fname)
        if fname.endswith(".docx"):
            all_text[fname] = read_docx(path)
        elif fname.endswith(".xlsx"):
            all_text[fname] = read_xlsx(path)
    return all_text

if __name__ == "__main__":
    docs = load_all_documents()
    for name, text in docs.items():
        print(f"{name}: {len(text)} chars")